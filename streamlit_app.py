import os
import time
import openai
import PyPDF2
import docx
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from dotenv import load_dotenv
import openpyxl  # For .xlsx files
import xlrd  # For .xls files
import subprocess  # For handling .doc files
import tempfile
import csv
from io import StringIO

# --- Configuration ---
load_dotenv()

# --- Security Note ---
# The COMPANY_NAME check provides basic access control but is not highly secure.
# Consider more robust methods if needed.
COMPANY_NAME = os.getenv("COMPANY_NAME")

# --- API Key Handling ---
# Prioritize reading from .env file. Ensure .env is in .gitignore!
# Use Streamlit secrets management if deploying on Streamlit Community Cloud.
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# --- Model & Pricing ---
# Update these based on the specific model and current OpenAI pricing
# Prices below are examples - PLEASE VERIFY CURRENT PRICING
# gpt-4o-mini pricing (example): $0.15 / 1M input, $0.60 / 1M output (as of May 2024)
PRICE_INPUT_PER_MILLION = 0.15  # Example price per million input tokens
PRICE_OUTPUT_PER_MILLION = 0.60 # Example price per million output tokens
AVG_TOKENS_PER_REQUEST = 780.5 # Your previous average - might need adjustment
# Output token estimation is difficult. Let's estimate based on a fixed average output length.
AVG_OUTPUT_TOKENS_PER_CLASSIFICATION = 50 # Estimate ~50 tokens for the CSV output line
DEFAULT_PROCESS_TIME_PER_FILE = 1.62
MAX_PROMPT_CHARS = 2000 # Max characters from document to send to OpenAI
LINES_PER_PAGE_ESTIMATE = 30 # For rough page count estimation from text lines

# --- File Handling ---
ROOT_DIR = "uploaded_files"
os.makedirs(ROOT_DIR, exist_ok=True)

# --- Streamlit App ---
st.set_page_config(page_title="Document Classification", layout="wide")
st.title("📄 Administrative Document Classification")

# --- Authentication ---
company_name_input = st.text_input("🔑 Enter your assigned access identifier", type="password")
if not company_name_input:
    st.info("Please enter the access identifier provided to you.")
    st.stop()
elif company_name_input != COMPANY_NAME:
    st.error("⚠️ Access Denied! Incorrect identifier.")
    st.stop()
else:
    st.success("✅ Access Granted")

# Check for API key from .env first, then ask if missing
if not OPENAI_API_KEY:
    st.warning("⚠️ OpenAI API Key not found in environment variables (.env).")
    OPENAI_API_KEY = st.text_input("🔑 Please enter your OpenAI API Key manually", type="password")
    if not OPENAI_API_KEY:
        st.warning("⚠️ An OpenAI API key is required to proceed.")
        st.stop()
    else:
         st.success("✅ API Key entered.")

try:
    openai.api_key = OPENAI_API_KEY
    # Optional: Perform a quick test call to validate the key early
    # openai.models.list()
    # st.success("✅ OpenAI API Key validated.")
except Exception as e:
    st.error(f"❌ Failed to initialize OpenAI client. Check your API key. Error: {e}")
    st.stop()


# --- Session State Initialization ---
if "classification_results" not in st.session_state:
    st.session_state.classification_results = []
if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()
if "error_files" not in st.session_state:
    st.session_state.error_files = {} # Store files that failed processing


# --- Helper Functions for Text Extraction ---

def estimate_pages_from_text(text):
    """Roughly estimate page count based on line breaks."""
    if not text:
        return 1
    return max(1, text.count("\n") // LINES_PER_PAGE_ESTIMATE + 1)

def estimate_pages_from_docx(doc):
    """Roughly estimate page count based on paragraph count."""
    return max(1, len(doc.paragraphs) // LINES_PER_PAGE_ESTIMATE + 1)

def extract_text_from_doc(file_path):
    """
    Extract text from old .doc format using various fallbacks.
    Requires external dependencies: python-docx (limited support), textract, antiword, LibreOffice.
    """
    text = ""
    pages = 1
    error_message = None

    # Attempt 1: python-docx (might work for some simple .doc files saved from newer Word)
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        pages = estimate_pages_from_docx(doc)
        if text.strip(): # Check if text was actually extracted
             st.info(f"Extracted .doc using python-docx (might be partial). Estimated pages: {pages}")
             return text, pages, None # Success
    except Exception as e:
        # st.warning(f"python-docx failed for .doc: {e}")
        pass # Continue to next method

    # Attempt 2: textract (if installed) - often requires antiword or other libs
    try:
        import textract
        text = textract.process(file_path).decode('utf-8', errors='replace')
        pages = estimate_pages_from_text(text)
        st.info(f"Extracted .doc using textract. Estimated pages: {pages}")
        return text, pages, None # Success
    except ImportError:
        error_message = "textract library not found. "
    except Exception as e:
        # st.warning(f"textract failed for .doc: {e}")
        error_message = f"textract error: {e}. "
        pass # Continue to next method

    # Attempt 3: antiword (if installed and in PATH)
    try:
        # Check if antiword exists
        subprocess.run(['antiword', '-V'], capture_output=True, check=True, text=True) # Check version to see if command exists
        # If it exists, run it
        result = subprocess.run(['antiword', file_path], capture_output=True, text=True, check=True)
        text = result.stdout
        pages = estimate_pages_from_text(text)
        st.info(f"Extracted .doc using antiword. Estimated pages: {pages}")
        return text, pages, None # Success
    except (FileNotFoundError, subprocess.CalledProcessError) as e:
        # st.warning(f"antiword failed or not found: {e}")
        error_message += "antiword command failed or not found. "
        pass # Continue to next method

    # Attempt 4: LibreOffice/OpenOffice (if installed and in PATH)
    try:
        # Check if soffice exists
        subprocess.run(['soffice', '--version'], capture_output=True, check=True, text=True)
        # Use temp dir for conversion output
        with tempfile.TemporaryDirectory() as temp_dir:
            subprocess.run(
                ['soffice', '--headless', '--convert-to', 'txt:Text', # Specify Text filter
                 '--outdir', temp_dir, file_path],
                check=True,
                timeout=30 # Add a timeout
            )
            base_name = os.path.basename(file_path)
            txt_filename = os.path.splitext(base_name)[0] + '.txt'
            converted_file_path = os.path.join(temp_dir, txt_filename)

            if os.path.exists(converted_file_path):
                with open(converted_file_path, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
                pages = estimate_pages_from_text(text)
                st.info(f"Extracted .doc using LibreOffice/soffice. Estimated pages: {pages}")
                return text, pages, None # Success
            else:
                 error_message += "soffice conversion did not produce expected .txt file. "
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
        # st.warning(f"LibreOffice/soffice failed or not found: {e}")
        error_message += f"LibreOffice/soffice command failed, not found, or timed out ({e}). "
        pass

    # Last resort: Failed extraction
    final_error = f"Could not extract text from .doc file '{os.path.basename(file_path)}'. Attempts failed. Requires 'textract', 'antiword', or 'LibreOffice'/'OpenOffice' in system PATH. Details: {error_message}"
    st.error(final_error)
    return "", 1, final_error # Return empty text, 1 page, and the error


def extract_text_from_xls(file_path):
    """Extract text from old .xls format using xlrd."""
    try:
        workbook = xlrd.open_workbook(file_path, logfile=open(os.devnull, 'w')) # Suppress xlrd logs
        sheet_texts = []
        for sheet_index in range(workbook.nsheets):
            sheet = workbook.sheet_by_index(sheet_index)
            sheet_text = []
            for row_idx in range(sheet.nrows):
                row_values = [str(sheet.cell_value(row_idx, col_idx)).strip()
                              for col_idx in range(sheet.ncols)
                              if str(sheet.cell_value(row_idx, col_idx)).strip()]
                if row_values:
                    sheet_text.append(" | ".join(row_values))
            if sheet_text: # Only add sheet if it has content
                sheet_texts.append(f"--- Sheet: {sheet.name} ---\n" + "\n".join(sheet_text))
        full_text = "\n\n".join(sheet_texts)
        return full_text, workbook.nsheets, None # Text, page/sheet count, error
    except Exception as e:
        error_msg = f"Error extracting text from .xls file '{os.path.basename(file_path)}': {str(e)}"
        st.error(error_msg)
        return "", 1, error_msg


def extract_text_from_xlsx(file_path):
    """Extract text from .xlsx format using openpyxl."""
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
        sheet_texts = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_text = []
            # Iterate rows with data
            for row in sheet.iter_rows(values_only=True):
                # Check if row has any non-None, non-empty string values
                row_values = [str(cell).strip() for cell in row if cell is not None]
                row_text = " | ".join(filter(None, row_values)) # Join only non-empty strings
                if row_text: # Only append if row has actual content
                    sheet_text.append(row_text)

            if sheet_text: # Only add sheet if it has content
                sheet_texts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(sheet_text))
        full_text = "\n\n".join(sheet_texts)
        return full_text, len(workbook.sheetnames), None # Text, page/sheet count, error
    except Exception as e:
        error_msg = f"Error extracting text from .xlsx file '{os.path.basename(file_path)}': {str(e)}"
        st.error(error_msg)
        return "", 1, error_msg


def extract_text(file_path):
    """Extracts text and estimates page count from various file types."""
    file_extension = os.path.splitext(file_path)[1].lower()
    text = ""
    pages = 1
    error = None

    try:
        if file_extension == ".pdf":
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file, strict=False) # Add strict=False for some malformed PDFs
                pages = len(reader.pages)
                extracted_pages = []
                for i, page in enumerate(reader.pages):
                    try:
                         page_text = page.extract_text()
                         if page_text:
                             extracted_pages.append(page_text)
                         # else:
                         #     st.warning(f"Page {i+1} in {os.path.basename(file_path)} yielded no text (possibly image-based).")
                    except Exception as page_error:
                        st.warning(f"Could not extract text from page {i+1} in {os.path.basename(file_path)}: {page_error}")
                text = "\n--- Page Break ---\n".join(extracted_pages)
            if not text.strip() and pages > 0:
                st.warning(f"PDF '{os.path.basename(file_path)}' ({pages} pages) yielded no extractable text. It might be image-only or scanned. OCR is needed for such files.")
                error = "PDF contains no extractable text (possibly image-based)."


        elif file_extension == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="replace") as file:
                text = file.read()
            pages = estimate_pages_from_text(text)

        elif file_extension == ".docx":
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            pages = estimate_pages_from_docx(doc) # Use rough estimate

        elif file_extension == ".doc":
            text, pages, error = extract_text_from_doc(file_path) # pages is estimate

        elif file_extension == ".xlsx":
            text, pages, error = extract_text_from_xlsx(file_path) # pages = sheet count

        elif file_extension == ".xls":
            text, pages, error = extract_text_from_xls(file_path) # pages = sheet count

        else:
            error = f"Unsupported file extension: {file_extension}"
            st.error(error)

    except Exception as e:
        error = f"Failed to process file '{os.path.basename(file_path)}'. Error: {str(e)}"
        st.error(error)
        text = "" # Ensure text is empty on error
        pages = 1

    return text.strip(), pages, error # Return text, pages, and any error message


def classify_document(text, file_name, pages):
    """Classifies document text using OpenAI API."""
    if not text:
        st.warning(f"Skipping classification for '{file_name}' due to empty extracted text.")
        return None # Indicate skipped classification

    prompt = f"""
    Analyze the following administrative document text and classify it.

    **Instructions:**
    1. Identify the main category (e.g., Report, Proposal, Decision, Invoice, Contract, Letter, Memo, Other).
    2. Identify the specific subcategory (e.g., Inspection Report, Investment Proposal, Research Evaluation, VAT Invoice, Service Contract, Official Letter). If multiple apply, separate with '/'. If none specific, use 'General'.
    3. Identify the relevant domain or industry (e.g., Government Administration, Science & Technology, Finance, Commerce, Legal, Education, Healthcare). If multiple apply, separate with '/'.
    4. Output **only** a single line of CSV data, using comma (,) as the delimiter. Do not add headers or any explanation.
    5. The output format must be exactly: File name,Main category,Subcategory,Domain/Industry,Pages
    6. Use Vietnamese for categories and domain.

    **Example Output:**
    1-C24MXD-00003085.pdf,Hóa đơn,Hóa đơn giá trị gia tăng,Thương mại/Dịch vụ,1

    **Document Details:**
    File Name: {file_name}
    Number of Pages/Sheets: {pages}

    **Document Text (first {MAX_PROMPT_CHARS} chars):**
    {text[:MAX_PROMPT_CHARS]}
    """

    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini", # Or your preferred model
            messages=[
                {"role": "system", "content": "You are an AI expert in classifying Vietnamese administrative documents into structured CSV format."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2, # Lower temperature for more deterministic output
            max_tokens=150 # Allow sufficient tokens for the output line
        )

        result_text = response.choices[0].message.content.strip()

        # --- Robust Parsing ---
        # Use StringIO and csv.reader for proper CSV handling
        csv_data = StringIO(result_text)
        reader = csv.reader(csv_data, delimiter=',')
        parsed_rows = list(reader)

        if not parsed_rows:
            st.warning(f"LLM returned empty response for '{file_name}'.")
            return None

        # Expecting only one row of data
        classification_data = parsed_rows[0]

        # Clean whitespace from each field
        classification_data = [field.strip() for field in classification_data]

        # Validate number of columns (expecting 5: File, Main, Sub, Domain, Pages)
        if len(classification_data) == 5:
            # Optional: Validate if the first column matches the expected filename (LLM might hallucinate)
            # if classification_data[0] != file_name:
            #     st.warning(f"LLM returned mismatching filename ('{classification_data[0]}') for '{file_name}'. Using original filename.")
            #     classification_data[0] = file_name

            # Optional: Try to validate/correct the page number if LLM gets it wrong
            try:
                llm_pages = int(classification_data[4])
                if llm_pages != pages:
                     st.warning(f"LLM page count ({llm_pages}) differs from extracted count ({pages}) for '{file_name}'. Using extracted count.")
                     classification_data[4] = str(pages) # Correct to original page count
            except ValueError:
                 st.warning(f"LLM returned non-integer page count ('{classification_data[4]}') for '{file_name}'. Using extracted count '{pages}'.")
                 classification_data[4] = str(pages) # Correct to original page count

            return classification_data # Return the validated list

        else:
            st.warning(f"LLM response for '{file_name}' had unexpected format (expected 5 CSV fields, got {len(classification_data)}): '{result_text}'")
            # Try to salvage if possible, or return None
            # Example: If it looks like it missed the filename, add it?
            if len(classification_data) == 4:
                 st.info("Attempting recovery by prepending filename.")
                 return [file_name] + classification_data[:3] + [str(pages)] # Add filename, take first 3 fields, add pages
            return None # Indicate failed parsing

    except openai.AuthenticationError:
        st.error("❌ OpenAI Authentication Error! Check your API Key.")
        st.stop() # Stop execution for auth errors

    except openai.OpenAIError as e:
        st.error(f"⚠️ OpenAI API Error for '{file_name}': {str(e)}")
        # Don't stop for individual file API errors, just log and continue
        return None # Indicate API error for this file

    except Exception as e:
        st.error(f"⚠️ Error parsing classification response for '{file_name}': {str(e)}. Response was: '{result_text}'")
        return None # Indicate parsing error


# --- File Uploader ---
st.header("1. Upload Files")
# Add note about dependencies
st.info("ℹ️ Processing `.doc` files may require external tools like `textract`, `antiword`, or `LibreOffice`/`OpenOffice` to be installed and accessible in your system's PATH.")

uploaded_files = st.file_uploader(
    "📂 Select files (.pdf, .txt, .docx, .doc, .xlsx, .xls)",
    type=["pdf", "txt", "docx", "doc", "xlsx", "xls"],
    accept_multiple_files=True,
    help="Upload one or more documents in the supported formats."
)

# --- Processing Section ---
if uploaded_files:
    st.header("2. Processing Options")
    total_files = len(uploaded_files)
    st.write(f"✅ **{total_files} file(s) selected.**")

    # Save uploaded files locally (handle potential errors)
    saved_file_paths = []
    save_errors = False
    with st.spinner("Saving uploaded files..."):
        for uploaded_file in uploaded_files:
            file_path = os.path.join(ROOT_DIR, uploaded_file.name)
            try:
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer()) # Use getbuffer() for efficiency
                saved_file_paths.append(file_path)
            except Exception as e:
                st.error(f"Error saving file '{uploaded_file.name}': {e}")
                save_errors = True

    if save_errors:
        st.error("Some files could not be saved. Please check permissions or disk space.")
        st.stop()
    elif not saved_file_paths:
         st.warning("No files were successfully saved.")
         st.stop()
    # else:
    #      st.success(f"Successfully saved {len(saved_file_paths)} files.")


    # --- Batching and Estimation ---
    col1, col2 = st.columns(2)
    with col1:
        batch_size = st.number_input(
            "🔢 Files per batch",
            min_value=1,
            max_value=total_files,
            value=min(10, total_files), # Default to 10 or total files if less
            step=1,
            help="Process files in smaller groups to manage resources or API rate limits."
        )
    with col2:
        delay_time = st.number_input(
            "⏳ Delay between batches (seconds)",
            min_value=0, # Allow 0 delay
            max_value=60,
            value=5,
            step=1,
            help="Pause between batches to avoid hitting API rate limits."
        )

    # Calculate estimations
    # Input cost is based on estimated total input tokens (chars / ~4 * files)
    # A better estimate might involve actually counting tokens after extraction if precision is critical
    estimated_total_input_tokens = total_files * AVG_TOKENS_PER_REQUEST # Using your average
    estimated_cost_input = (estimated_total_input_tokens / 1_000_000) * PRICE_INPUT_PER_MILLION

    # Output cost based on average classification tokens per file
    estimated_total_output_tokens = total_files * AVG_OUTPUT_TOKENS_PER_CLASSIFICATION
    estimated_cost_output = (estimated_total_output_tokens / 1_000_000) * PRICE_OUTPUT_PER_MILLION

    estimated_total_cost = estimated_cost_input + estimated_cost_output

    st.subheader("📊 Estimated Cost & Time")
    st.info(f"""
    *   **Input Tokens (Est.):** {estimated_total_input_tokens:,.0f}
    *   **Output Tokens (Est.):** {estimated_total_output_tokens:,.0f}
    *   💰 **Estimated Total OpenAI Cost:** ${estimated_total_cost:.4f} (Based on example pricing for gpt-4o-mini. Actual cost may vary.)
    """)

    # Time estimation
    estimated_time_processing_only = total_files * DEFAULT_PROCESS_TIME_PER_FILE
    num_batches = (total_files + batch_size - 1) // batch_size # Ceiling division
    total_delay_time = max(0, num_batches - 1) * delay_time if process_batches else 0
    estimated_time_with_delay = estimated_time_processing_only + total_delay_time

    st.info(f"""
    *   ⏳ **Est. Processing Time (excl. delay):** {estimated_time_processing_only:.1f} sec (~{(estimated_time_processing_only / 60):.1f} min)
    *   ⏳ **Est. Total Time (incl. potential delay):** {estimated_time_with_delay:.1f} sec (~{(estimated_time_with_delay / 60):.1f} min)
    """)
    st.caption("Actual processing time depends on file complexity, API response times, and system load.")


    # --- Processing Execution ---
    st.subheader("🚀 Start Processing")
    process_button_cols = st.columns(2)
    # Note: Streamlit reruns script on button press. We use session state to track progress.
    start_processing = process_button_cols[0].button(f"🔄 Process {total_files} File(s) Now", type="primary")

    if start_processing:
        st.session_state.processing_started = True
        st.session_state.current_file_index = 0 # Track overall progress
        st.session_state.error_files = {} # Reset errors for new run
        # Clear previous results *if* starting a new full run
        st.session_state.classification_results = []
        st.session_state.processed_files = set()


    # --- Progress Display and Processing Loop ---
    if st.session_state.get("processing_started", False):
        progress_bar = st.progress(0.0, text="Starting processing...")
        files_to_process = [fp for fp in saved_file_paths if os.path.basename(fp) not in st.session_state.processed_files]
        processed_count_this_run = 0

        if not files_to_process:
             st.info("All selected files have already been processed in this session.")
             st.session_state.processing_started = False # Stop processing state
        else:
            total_to_process_now = len(files_to_process)
            st.write(f"Processing {total_to_process_now} remaining files...")

            file_batches = [files_to_process[i:i + batch_size] for i in range(0, total_to_process_now, batch_size)]

            for batch_index, batch in enumerate(file_batches):
                st.write(f"--- Processing Batch {batch_index + 1} / {len(file_batches)} ---")
                batch_start_time = time.time()

                for file_index_in_batch, file_path in enumerate(batch):
                    file_name = os.path.basename(file_path)
                    overall_file_number = st.session_state.current_file_index + 1

                    progress_text = f"Processing file {overall_file_number}/{total_files}: {file_name}"
                    progress_bar.progress(st.session_state.current_file_index / total_files, text=progress_text)
                    # st.write(f"🔄 {progress_text}") # Can be verbose, use progress bar text

                    # 1. Extract Text
                    with st.spinner(f"Extracting text from {file_name}..."):
                         text, pages, extract_error = extract_text(file_path)

                    if extract_error:
                        st.error(f"Extraction failed for {file_name}: {extract_error}")
                        st.session_state.error_files[file_name] = f"Extraction Error: {extract_error}"
                        st.session_state.processed_files.add(file_name) # Mark as processed (with error)
                        st.session_state.current_file_index += 1
                        continue # Skip to next file

                    if not text:
                        st.warning(f"No text extracted from {file_name}. Skipping classification.")
                        st.session_state.error_files[file_name] = "Extraction Error: No text content found/extracted."
                        st.session_state.processed_files.add(file_name) # Mark as processed (with error)
                        st.session_state.current_file_index += 1
                        continue # Skip to next file

                    # 2. Classify Document
                    with st.spinner(f"Classifying {file_name} ({pages} pages/sheets)..."):
                         classification_result = classify_document(text, file_name, pages)

                    if classification_result:
                        # Append correctly parsed result
                        st.session_state.classification_results.append(classification_result)
                        st.session_state.processed_files.add(file_name) # Mark as successfully processed
                        # st.write(f"✅ Classified: {classification_result[1]}, {classification_result[2]}, {classification_result[3]}")
                    else:
                        # Handle classification failure (already logged in function)
                        st.session_state.error_files[file_name] = "Classification Error (check logs above)"
                        st.session_state.processed_files.add(file_name) # Mark as processed (with error)

                    processed_count_this_run += 1
                    st.session_state.current_file_index += 1


                # Batch finished
                batch_end_time = time.time()
                st.write(f"--- Batch {batch_index + 1} finished in {batch_end_time - batch_start_time:.2f} seconds ---")

                # Delay if not the last batch
                if delay_time > 0 and batch_index < len(file_batches) - 1:
                    st.write(f"⏳ Waiting {delay_time} sec before next batch...")
                    time.sleep(delay_time)

            # Update progress bar to 100% at the end
            progress_bar.progress(1.0, text=f"Processing complete. Processed {processed_count_this_run} files in this run.")
            st.success(f"✅ Processing finished!")
            st.session_state.processing_started = False # Reset processing state


# --- Display Results ---
if st.session_state.classification_results:
    st.header("3. Classification Results")

    df = pd.DataFrame(
        st.session_state.classification_results,
        columns=["File name", "Main category", "Subcategory", "Domain/Industry", "Pages"]
    )

    # Display DataFrame
    st.dataframe(df, use_container_width=True)
    st.info(f"Total classified documents: {len(df)}")
    st.caption("Note: Page counts for non-PDF files are estimates (often sheet counts for Excel).")


    # --- Display Errors ---
    if st.session_state.error_files:
         st.subheader("⚠️ Processing Errors")
         error_df = pd.DataFrame(list(st.session_state.error_files.items()), columns=["File Name", "Error"])
         st.dataframe(error_df, use_container_width=True)


    # --- Visualization ---
    st.header("4. Visualize Results")
    if not df.empty:
        col1, col2 = st.columns([1, 2]) # Make chart options column smaller

        with col1:
            chart_type = st.selectbox("📊 Chart Type", ["Bar", "Pie"], key="chart_type")
            # Filter out potential empty/placeholder categories before counting
            valid_categories = df["Main category"].dropna().unique()
            valid_subcategories = df["Subcategory"].dropna().unique()
            valid_domains = df["Domain/Industry"].dropna().unique()

            metric_options = []
            if len(valid_categories) > 0: metric_options.append("Main category")
            if len(valid_subcategories) > 0: metric_options.append("Subcategory")
            if len(valid_domains) > 0: metric_options.append("Domain/Industry")

            if not metric_options:
                st.warning("No valid data available for plotting.")
            else:
                 metric = st.selectbox("📌 Metric to Visualize", metric_options, key="metric_select")

                 # Filter DataFrame for the chosen metric to avoid errors with NaN in plots
                 plot_df = df[[metric]].dropna()


        with col2:
             if metric_options and metric and not plot_df.empty:
                 st.subheader(f"{chart_type} Chart: Distribution by {metric}")
                 plt.figure(figsize=(10, 6)) # Adjusted figure size

                 if chart_type == "Bar":
                     # Consider plotting only top N categories if there are too many
                     top_n = 20
                     value_counts = plot_df[metric].value_counts()
                     top_values = value_counts.nlargest(top_n).index
                     plot_subset = plot_df[plot_df[metric].isin(top_values)]

                     sns.countplot(data=plot_subset, y=metric, order=top_values, palette="viridis") # Use y-axis for better label readability
                     plt.title(f'Top {min(top_n, len(value_counts))} {metric} Distribution')
                     plt.xlabel("Number of Documents")
                     plt.ylabel(metric)
                     plt.tight_layout() # Adjust layout
                     st.pyplot(plt)
                     if len(value_counts) > top_n:
                         st.caption(f"Showing top {top_n} categories. Total unique: {len(value_counts)}")


                 elif chart_type == "Pie":
                     # Pie charts are less effective for many categories, consider top N
                     top_n = 10
                     value_counts = plot_df[metric].value_counts()
                     if len(value_counts) > top_n:
                         top_n_counts = value_counts.nlargest(top_n)
                         other_count = value_counts.iloc[top_n:].sum()
                         if other_count > 0:
                            top_n_counts['Other'] = other_count
                         data_to_plot = top_n_counts
                         title = f'Top {top_n} {metric} Distribution (including Other)'
                     else:
                         data_to_plot = value_counts
                         title = f'{metric} Distribution'

                     if not data_to_plot.empty:
                          data_to_plot.plot.pie(autopct="%1.1f%%", startangle=90, counterclock=False,
                                                  wedgeprops=dict(width=0.4), pctdistance=0.8, # Doughnut style
                                                  figsize=(8, 8)) # Make pie chart larger
                          plt.title(title)
                          plt.ylabel('') # Hide y-label for pie chart
                          st.pyplot(plt)
                     else:
                          st.warning("No data to plot for the selected metric.")


             elif not metric_options:
                  st.warning("No categories found to visualize.")
             else:
                  st.warning(f"No valid data found for metric: {metric}")

    else:
        st.warning("No results available to visualize.")


    # --- Download ---
    st.header("5. Download Results")
    try:
        csv_output = df.to_csv(index=False, encoding='utf-8-sig') # utf-8-sig for Excel compatibility
        st.download_button(
            label="📥 Download Classification Results (CSV)",
            data=csv_output,
            file_name="document_classification_results.csv",
            mime="text/csv",
        )
    except Exception as e:
        st.error(f"Failed to generate CSV for download: {e}")


# --- Clear State Option ---
st.sidebar.header("Options")
if st.sidebar.button("Clear Results & State"):
    st.session_state.classification_results = []
    st.session_state.processed_files = set()
    st.session_state.error_files = {}
    if 'processing_started' in st.session_state:
        del st.session_state['processing_started']
    if 'current_file_index' in st.session_state:
        del st.session_state['current_file_index']
    # Optionally clear uploaded files too? Be careful with this.
    # Could remove files from ROOT_DIR if desired, but might be unexpected.
    st.success("Cleared current session results and errors.")
    st.rerun() # Rerun to reflect cleared state